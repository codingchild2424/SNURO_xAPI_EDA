#word 자동화 코드

from docx import Document

def make_LA_report(picture1, picture2, picture3):
    document = Document()

    document.add_heading('주간 학습분석 결과 보고서', level = 0)

    document.add_paragraph('본 문서는 1주일 간 학습자들의 행동 데이터를 수집하여 분석한 보고서입니다.')
    document.add_paragraph('SNURO멘토님들은 본 문서를 확인하시고, 교육적 의사결정을 내려주시길 바랍니다.')

    document.add_paragraph('1. 활동량 분석')
    document.add_picture(picture1)

    document.add_paragraph('2. 활동량 분석')
    document.add_picture(picture2)

    document.add_paragraph('1. 활동량 분석')
    document.add_picture(picture3)

    document.save("output.docx")